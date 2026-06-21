from typing import Dict, Any, List
from docx import Document


def parse_docx(file_path: str) -> Dict[str, Any]:
    document = Document(file_path)

    paragraphs: List[str] = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    tables = []
    for table in document.tables:
        table_rows = []
        for row in table.rows:
            table_rows.append([cell.text.strip() for cell in row.cells])
        tables.append(table_rows)

    table_text_parts = []
    for table_index, table in enumerate(tables, start=1):
        table_text_parts.append(f"Table {table_index}:")
        for row in table:
            table_text_parts.append(" | ".join(row))

    combined_text = "\n".join(paragraphs + table_text_parts)

    return {
        "content_type": "document",
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "paragraphs": paragraphs,
        "tables": tables,
        "text": combined_text,
    }